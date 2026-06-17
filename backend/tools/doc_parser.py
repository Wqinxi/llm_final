import base64
import io
import os
import tempfile
from pdfplumber import open as pdf_open
from docx import Document
from PyPDF2 import PdfReader


def decode_base64(base64_str: str) -> bytes:
    """base64字符串转二进制"""
    if base64_str.startswith("data:"):
        base64_str = base64_str.split(",")[-1]
    return base64.b64decode(base64_str)


def parse_document(file_name: str, file_base64: str) -> str:
    """
    解析 txt/md/pdf/doc/docx 文档
    支持 Windows 下使用 pywin32 + Microsoft Word 解析老式 .doc 格式
    """
    try:
        file_bytes = decode_base64(file_base64)
        lower_name = file_name.lower()
        full_text = ""

        if lower_name.endswith((".txt", ".md")):
            # 多编码兼容
            buf = io.BytesIO(file_bytes)
            for enc in ["utf-8", "gbk", "gb2312"]:
                try:
                    full_text = buf.read().decode(enc)
                    break
                except Exception:
                    buf.seek(0)
                    continue
            else:
                full_text = buf.read().decode("utf-8", errors="ignore")

        elif lower_name.endswith(".pdf"):
            buf = io.BytesIO(file_bytes)
            try:
                with pdf_open(buf) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        full_text += page_text + "\n"
            except Exception:
                buf.seek(0)
                reader = PdfReader(buf)
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    full_text += page_text + "\n"

        elif lower_name.endswith(".docx"):
            buf = io.BytesIO(file_bytes)
            doc = Document(buf)
            # 解析段落
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            # 解析表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join([cell.text for cell in row.cells])
                    full_text += row_text + "\n"

        elif lower_name.endswith(".doc"):
            # 老式 .doc 格式：临时写入本地文件，通过 Word COM 解析
            temp_file = None
            try:
                # 创建临时 .doc 文件
                with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
                    f.write(file_bytes)
                    temp_file = f.name

                # 调用 pywin32 + Word 读取
                import win32com.client as wc
                word = wc.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(temp_file))
                full_text = doc.Content.Text
                doc.Close(False)
                word.Quit()

            except ImportError:
                # 未安装 pywin32
                full_text = f"【提示】解析 {file_name} 需要安装 pywin32，执行：pip install pywin32"
            except Exception:
                # 无 Word 或解析失败（非Windows环境/软件异常）
                full_text = f"【提示】当前环境无法使用Word解析 {file_name}，建议另存为 .docx 格式后上传"
            finally:
                # 强制删除临时文件，防止残留
                if temp_file and os.path.exists(temp_file):
                    try:
                        os.unlink(temp_file)
                    except Exception:
                        pass

        else:
            full_text = f"【不支持格式】无法解析文件：{file_name}"

        # 清理多余空行
        lines = [line.strip() for line in full_text.splitlines() if line.strip()]
        clean_text = "\n".join(lines)
        return f"【上传文档：{file_name}】\n{clean_text}\n"

    except Exception as e:
        return f"【文档解析失败：{file_name}】\n错误详情：{str(e)}"